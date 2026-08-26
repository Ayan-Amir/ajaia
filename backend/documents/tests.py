import io

from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from docx import Document as DocxDocument
from rest_framework import status
from rest_framework.test import APITestCase

from documents.choices import DocumentRole, SharePermission
from documents.models import Document, DocumentShare

FORBIDDEN = "You do not have permission to perform this action."
UNAUTHENTICATED = "Authentication credentials were not provided."


def build_docx(heading_level=1):
    """Build a real .docx in memory so the import path is exercised for real."""
    document = DocxDocument()
    document.add_heading("Quarterly Report", level=heading_level)

    paragraph = document.add_paragraph("Revenue was ")
    paragraph.add_run("bold").bold = True
    paragraph.add_run(" and ")
    paragraph.add_run("italic").italic = True

    document.add_paragraph("First point", style="List Bullet")
    document.add_paragraph("Second point", style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class DocumentAccessAPIViewTestCase(APITestCase):
    @classmethod
    def setUpTestData(cls):
        cls.owner = User.objects.create_user("owner")
        cls.viewer = User.objects.create_user("viewer")
        cls.editor = User.objects.create_user("editor")
        cls.stranger = User.objects.create_user("stranger")

        cls.document = Document.objects.create(
            owner=cls.owner, title="Spec", content="<p>Original</p>"
        )
        DocumentShare.objects.create(
            document=cls.document,
            shared_with=cls.viewer,
            permission=SharePermission.VIEW,
        )
        DocumentShare.objects.create(
            document=cls.document,
            shared_with=cls.editor,
            permission=SharePermission.EDIT,
        )
        cls.detail_url = reverse("document-detail", args=[cls.document.pk])
        cls.list_url = reverse("document-list")

    def test_unauthenticated_request_is_rejected(self):
        response = self.client.get(self.detail_url)
        self.assertEqual(response.status_code, status.HTTP_401_UNAUTHORIZED)
        self.assertEqual(response.data["detail"], UNAUTHENTICATED)

    def test_owner_can_retrieve_document(self):
        self.client.force_authenticate(self.owner)
        response = self.client.get(self.detail_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data["role"], DocumentRole.OWNER)
        self.assertEqual(response.data["content"], "<p>Original</p>")

    def test_stranger_cannot_see_document(self):
        self.client.force_authenticate(self.stranger)
        self.assertEqual(
            self.client.get(self.detail_url).status_code, status.HTTP_404_NOT_FOUND
        )

    def test_list_separates_owned_from_shared(self):
        self.client.force_authenticate(self.viewer)
        response = self.client.get(self.list_url)
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual([doc["role"] for doc in response.data], [DocumentRole.VIEWER])

    def test_stranger_list_is_empty(self):
        self.client.force_authenticate(self.stranger)
        self.assertEqual(self.client.get(self.list_url).data, [])

    def test_viewer_can_read_but_not_edit(self):
        self.client.force_authenticate(self.viewer)
        self.assertEqual(
            self.client.get(self.detail_url).status_code, status.HTTP_200_OK
        )

        response = self.client.patch(self.detail_url, {"content": "<p>Nope</p>"})
        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)
        self.assertEqual(response.data["detail"], FORBIDDEN)

        self.document.refresh_from_db()
        self.assertEqual(self.document.content, "<p>Original</p>")

    def test_editor_can_update_content(self):
        self.client.force_authenticate(self.editor)
        response = self.client.patch(self.detail_url, {"content": "<p>Edited</p>"})
        self.assertEqual(response.status_code, status.HTTP_200_OK)

        self.document.refresh_from_db()
        self.assertEqual(self.document.content, "<p>Edited</p>")

    def test_editor_cannot_delete_document(self):
        self.client.force_authenticate(self.editor)
        response = self.client.delete(self.detail_url)
        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)
        self.assertTrue(Document.objects.filter(pk=self.document.pk).exists())

    def test_owner_can_rename_document(self):
        self.client.force_authenticate(self.owner)
        response = self.client.patch(self.detail_url, {"title": "Renamed"})
        self.assertEqual(response.status_code, status.HTTP_200_OK)

        self.document.refresh_from_db()
        self.assertEqual(self.document.title, "Renamed")

    def test_blank_title_is_rejected(self):
        self.client.force_authenticate(self.owner)
        response = self.client.patch(self.detail_url, {"title": "   "})
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)

    def test_script_tags_are_stripped_from_content(self):
        self.client.force_authenticate(self.owner)
        self.client.patch(
            self.detail_url,
            {"content": "<p>Safe</p><script>alert('xss')</script>"},
            format="json",
        )
        self.document.refresh_from_db()
        self.assertEqual(self.document.content, "<p>Safe</p>")


class DocumentShareAPIViewTestCase(APITestCase):
    @classmethod
    def setUpTestData(cls):
        cls.owner = User.objects.create_user("owner")
        cls.editor = User.objects.create_user("editor")
        cls.target = User.objects.create_user("target")
        cls.document = Document.objects.create(owner=cls.owner, title="Plan")
        DocumentShare.objects.create(
            document=cls.document,
            shared_with=cls.editor,
            permission=SharePermission.EDIT,
        )
        cls.share_url = reverse("document-share", args=[cls.document.pk])

    def test_owner_can_share_document(self):
        self.client.force_authenticate(self.owner)
        response = self.client.post(
            self.share_url,
            {"user_id": self.target.pk, "permission": SharePermission.EDIT},
        )
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertEqual(response.data["shared_with"]["username"], "target")

    def test_resharing_updates_permission_instead_of_duplicating(self):
        self.client.force_authenticate(self.owner)
        for permission in (SharePermission.VIEW, SharePermission.EDIT):
            self.client.post(
                self.share_url, {"user_id": self.target.pk, "permission": permission}
            )

        shares = DocumentShare.objects.filter(
            document=self.document, shared_with=self.target
        )
        self.assertEqual(shares.count(), 1)
        self.assertEqual(shares.first().permission, SharePermission.EDIT)

    def test_owner_cannot_share_with_themselves(self):
        self.client.force_authenticate(self.owner)
        response = self.client.post(self.share_url, {"user_id": self.owner.pk})
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)

    def test_editor_cannot_share_document(self):
        self.client.force_authenticate(self.editor)
        response = self.client.post(self.share_url, {"user_id": self.target.pk})
        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)
        self.assertEqual(response.data["detail"], FORBIDDEN)

    def test_owner_can_revoke_share(self):
        share = DocumentShare.objects.create(
            document=self.document,
            shared_with=self.target,
            permission=SharePermission.VIEW,
        )
        self.client.force_authenticate(self.owner)
        response = self.client.delete(
            reverse("document-share-detail", args=[self.document.pk, share.pk])
        )
        self.assertEqual(response.status_code, status.HTTP_204_NO_CONTENT)
        self.assertFalse(DocumentShare.objects.filter(pk=share.pk).exists())

    def test_revoked_user_loses_access(self):
        share = DocumentShare.objects.create(
            document=self.document,
            shared_with=self.target,
            permission=SharePermission.VIEW,
        )
        detail_url = reverse("document-detail", args=[self.document.pk])

        self.client.force_authenticate(self.target)
        self.assertEqual(self.client.get(detail_url).status_code, status.HTTP_200_OK)

        share.delete()
        self.assertEqual(
            self.client.get(detail_url).status_code, status.HTTP_404_NOT_FOUND
        )


class DocumentImportAPIViewTestCase(APITestCase):
    @classmethod
    def setUpTestData(cls):
        cls.user = User.objects.create_user("importer")
        cls.url = reverse("document-import")

    def setUp(self):
        self.client.force_authenticate(self.user)

    def _upload(self, name, body):
        return self.client.post(
            self.url,
            {"file": SimpleUploadedFile(name, body.encode("utf-8"))},
            format="multipart",
        )

    def test_txt_import_creates_owned_document(self):
        response = self._upload("meeting-notes.txt", "First para\n\nSecond para")
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertEqual(response.data["title"], "meeting-notes")
        self.assertEqual(response.data["content"], "<p>First para</p><p>Second para</p>")

        document = Document.objects.get(pk=response.data["id"])
        self.assertEqual(document.owner, self.user)

    def test_docx_import_preserves_formatting(self):
        response = self.client.post(
            self.url,
            {"file": SimpleUploadedFile("quarterly-report.docx", build_docx())},
            format="multipart",
        )
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertEqual(response.data["title"], "quarterly-report")

        content = response.data["content"]
        self.assertIn("Quarterly Report", content)
        self.assertIn("<strong>bold</strong>", content)
        self.assertIn("<em>italic</em>", content)
        self.assertIn("<li>", content)
        self.assertEqual(Document.objects.get(pk=response.data["id"]).owner, self.user)

    def test_docx_deep_headings_degrade_to_h3(self):
        response = self.client.post(
            self.url,
            {"file": SimpleUploadedFile("deep.docx", build_docx(heading_level=5))},
            format="multipart",
        )
        content = response.data["content"]
        self.assertIn("<h3>Quarterly Report</h3>", content)
        self.assertNotIn("<h5", content)

    def test_corrupt_docx_is_rejected(self):
        response = self.client.post(
            self.url,
            {"file": SimpleUploadedFile("broken.docx", b"this is not a real docx")},
            format="multipart",
        )
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("could not be read", response.data["file"][0])
        self.assertFalse(Document.objects.exists())

    def test_markdown_import_preserves_formatting(self):
        response = self._upload("spec.md", "# Title\n\nSome **bold** text\n")
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn("<h1>Title</h1>", response.data["content"])
        self.assertIn("<strong>bold</strong>", response.data["content"])

    def test_txt_import_escapes_html(self):
        response = self._upload("raw.txt", "<script>alert(1)</script>")
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertNotIn("<script>", response.data["content"])

    def test_unsupported_extension_is_rejected(self):
        response = self._upload("resume.pdf", "not really a pdf")
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("Only .txt, .md and .docx", response.data["file"][0])
        self.assertFalse(Document.objects.exists())

    def test_import_requires_authentication(self):
        self.client.force_authenticate(None)
        response = self._upload("notes.txt", "hello")
        self.assertEqual(response.status_code, status.HTTP_401_UNAUTHORIZED)
